"""
Test suite: Kiểm tra parse_document xử lý được từng định dạng file nào.
Tự động tạo file mẫu cho mỗi format rồi gọi parser.
"""

import os
import sys
import tempfile

# Ensure project root is on path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from src.utils.parser import parse_document

SAMPLE_TEXT_VI = (
    "CLO1: Giải thích được các khái niệm cơ bản về trí tuệ nhân tạo.\n"
    "CLO2: Áp dụng thuật toán tìm kiếm trên đồ thị để giải quyết bài toán thực tế.\n"
    "CLO3: Thiết kế một hệ thống AI đơn giản cho bài toán phân loại."
)


SAMPLE_TEXT_PDF = (
    "CLO1: Explain fundamental concepts of artificial intelligence.\n"
    "CLO2: Apply graph search algorithms to solve real-world problems.\n"
    "CLO3: Design a simple AI system for classification tasks."
)


def _make_pdf(path: str, text: str) -> None:
    """Tạo file PDF text-based bằng fpdf2 (nếu có) hoặc reportlab."""
    try:
        from fpdf import FPDF
        from fpdf.enums import XPos, YPos

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        # Dùng text ASCII để tránh lỗi encoding với Helvetica
        for line in SAMPLE_TEXT_PDF.split("\n"):
            pdf.cell(0, 10, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.output(path)
        return
    except ImportError:
        pass

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        c = canvas.Canvas(path, pagesize=A4)
        y = 750
        for line in text.split("\n"):
            c.drawString(50, y, line)
            y -= 20
        c.save()
        return
    except ImportError:
        pass

    # Fallback: tạo PDF thủ công (minimal valid PDF)
    content = text.replace("\n", " ")
    pdf_bytes = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
    )
    stream = f"BT /F1 12 Tf 50 750 Td ({content}) Tj ET"
    stream_bytes = stream.encode("latin-1", errors="replace")
    pdf_bytes += f"4 0 obj<</Length {len(stream_bytes)}>>\nstream\n".encode() + stream_bytes + b"\nendstream\nendobj\n"
    xref_offset = len(pdf_bytes)
    pdf_bytes += (
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"0000000306 00000 n \n"
        b"0000000250 00000 n \n"
        b"trailer<</Size 6/Root 1 0 R>>\n"
        b"startxref\n" + str(xref_offset).encode() + b"\n%%EOF"
    )
    with open(path, "wb") as f:
        f.write(pdf_bytes)


def _make_docx(path: str, text: str) -> None:
    """Tạo file DOCX sử dụng python-docx."""
    import docx

    doc = docx.Document()
    doc.add_heading("Syllabus - Trí Tuệ Nhân Tạo", level=1)
    for line in text.split("\n"):
        doc.add_paragraph(line)
    # Thêm bảng để test table extraction
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Tuần"
    table.cell(0, 1).text = "Nội dung"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "Giới thiệu AI"
    doc.save(path)


def _make_txt(path: str, text: str) -> None:
    """Tạo file TXT thuần."""
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)


def _make_txt_latin1(path: str) -> None:
    """Tạo file TXT với encoding latin-1 để test fallback."""
    content = "CLO1: Explain basic concepts.\nCLO2: Apply search algorithms."
    with open(path, "w", encoding="latin-1") as f:
        f.write(content)


def _make_fake_doc(path: str, text: str) -> None:
    """Tạo file .doc giả (thực tế là binary ngẫu nhiên, không phải OLE format)."""
    with open(path, "wb") as f:
        f.write(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")  # OLE magic bytes
        f.write(text.encode("utf-8", errors="replace"))


def _make_unsupported(path: str) -> None:
    """Tạo file với extension không hỗ trợ."""
    with open(path, "w") as f:
        f.write("This should not be parseable")


# ──────────────────────────────────────────────────────────────
# Test functions
# ──────────────────────────────────────────────────────────────


def test_pdf_parsing():
    """Test 1: PDF text-based → phải trích xuất được text."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        path = tmp.name
    try:
        _make_pdf(path, SAMPLE_TEXT_VI)
        result = parse_document(path)
        print(f"\n{'=' * 60}")
        print("TEST 1: PDF (.pdf)")
        print(f"{'=' * 60}")
        print(f"  File size : {os.path.getsize(path)} bytes")
        print(f"  Extracted : {len(result)} chars")
        print(f"  Preview   : {result[:120]}...")
        if result.strip():
            print("  ✅ PASS — Trích xuất text từ PDF thành công")
        else:
            print("  ❌ FAIL — Không trích xuất được text từ PDF")
        return bool(result.strip())
    finally:
        os.unlink(path)


def test_docx_parsing():
    """Test 2: DOCX → phải trích xuất paragraphs + tables."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        path = tmp.name
    try:
        _make_docx(path, SAMPLE_TEXT_VI)
        result = parse_document(path)
        print(f"\n{'=' * 60}")
        print("TEST 2: DOCX (.docx)")
        print(f"{'=' * 60}")
        print(f"  File size : {os.path.getsize(path)} bytes")
        print(f"  Extracted : {len(result)} chars")
        print(f"  Preview   : {result[:120]}...")
        has_paragraph = "CLO1" in result
        has_table = "Giới thiệu AI" in result or "Nội dung" in result
        print(f"  Paragraphs: {'✅' if has_paragraph else '❌'}")
        print(f"  Tables    : {'✅' if has_table else '❌'}")
        if has_paragraph and has_table:
            print("  ✅ PASS — Trích xuất cả paragraph và table từ DOCX")
        elif has_paragraph:
            print("  ⚠️ PARTIAL — Chỉ trích xuất paragraph, thiếu table")
        else:
            print("  ❌ FAIL — Không trích xuất được text từ DOCX")
        return has_paragraph
    finally:
        os.unlink(path)


def test_txt_utf8_parsing():
    """Test 3: TXT UTF-8 → đọc trực tiếp."""
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w") as tmp:
        path = tmp.name
    try:
        _make_txt(path, SAMPLE_TEXT_VI)
        result = parse_document(path)
        print(f"\n{'=' * 60}")
        print("TEST 3: TXT UTF-8 (.txt)")
        print(f"{'=' * 60}")
        print(f"  File size : {os.path.getsize(path)} bytes")
        print(f"  Extracted : {len(result)} chars")
        print(f"  Preview   : {result[:120]}...")
        if "CLO1" in result and "CLO3" in result:
            print("  ✅ PASS — Đọc file TXT UTF-8 thành công")
        else:
            print("  ❌ FAIL — Không đọc được file TXT UTF-8")
        return "CLO1" in result
    finally:
        os.unlink(path)


def test_txt_latin1_parsing():
    """Test 4: TXT Latin-1 → fallback encoding."""
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
        path = tmp.name
    try:
        _make_txt_latin1(path)
        result = parse_document(path)
        print(f"\n{'=' * 60}")
        print("TEST 4: TXT Latin-1 Fallback (.txt)")
        print(f"{'=' * 60}")
        print(f"  File size : {os.path.getsize(path)} bytes")
        print(f"  Extracted : {len(result)} chars")
        print(f"  Preview   : {result[:120]}...")
        if "CLO1" in result:
            print("  ✅ PASS — Fallback Latin-1 hoạt động")
        else:
            print("  ❌ FAIL — Fallback Latin-1 không hoạt động")
        return "CLO1" in result
    finally:
        os.unlink(path)


def test_doc_old_format():
    """Test 5: .doc (Word cũ) → python-docx không hỗ trợ, expected FAIL."""
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        path = tmp.name
    try:
        _make_fake_doc(path, SAMPLE_TEXT_VI)
        result = parse_document(path)
        print(f"\n{'=' * 60}")
        print("TEST 5: DOC cũ (.doc) — Expected: FAIL")
        print(f"{'=' * 60}")
        print(f"  File size : {os.path.getsize(path)} bytes")
        print(f"  Extracted : {len(result)} chars")
        if result.strip():
            print(f"  Preview   : {result[:120]}...")
            print("  ⚠️ UNEXPECTED PASS — python-docx parse được .doc?")
        else:
            print("  ✅ CONFIRMED — .doc cũ KHÔNG được hỗ trợ (trả về empty)")
            print("  → python-docx chỉ hỗ trợ .docx, KHÔNG hỗ trợ .doc cũ (OLE format)")
        return False  # Expected fail
    finally:
        os.unlink(path)


def test_unsupported_extension():
    """Test 6: Extension không hỗ trợ (.xlsx, .pptx, v.v.) → expected empty."""
    results = {}
    for ext in [".xlsx", ".pptx", ".csv", ".html", ".md"]:
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            path = tmp.name
        try:
            with open(path, "w") as f:
                f.write("test content")
            result = parse_document(path)
            results[ext] = bool(result.strip())
        finally:
            os.unlink(path)

    print(f"\n{'=' * 60}")
    print("TEST 6: Unsupported Extensions")
    print(f"{'=' * 60}")
    all_rejected = True
    for ext, parsed in results.items():
        status = "❌ PARSED (should reject)" if parsed else "✅ Rejected"
        print(f"  {ext:8s} : {status}")
        if parsed:
            all_rejected = False
    if all_rejected:
        print("  ✅ PASS — Tất cả extension không hỗ trợ đều bị reject")
    else:
        print("  ❌ FAIL — Một số extension không hỗ trợ vẫn được parse")
    return all_rejected


def test_empty_file():
    """Test 7: File rỗng → phải trả về empty string."""
    results = {}
    for ext in [".pdf", ".docx", ".txt"]:
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            path = tmp.name
        try:
            # File rỗng (0 bytes cho txt, hoặc vài bytes invalid cho pdf/docx)
            if ext == ".txt":
                with open(path, "w") as f:
                    f.write("")
            # pdf và docx đã là empty file từ NamedTemporaryFile
            result = parse_document(path)
            results[ext] = result
        finally:
            os.unlink(path)

    print(f"\n{'=' * 60}")
    print("TEST 7: Empty / Invalid Files")
    print(f"{'=' * 60}")
    for ext, result in results.items():
        chars = len(result) if result else 0
        status = "✅ Empty result" if chars == 0 else f"⚠️ Got {chars} chars"
        print(f"  {ext:8s} : {status}")


def test_nonexistent_file():
    """Test 8: File không tồn tại → phải trả về empty string."""
    result = parse_document("/nonexistent/path/fake_syllabus.pdf")
    print(f"\n{'=' * 60}")
    print("TEST 8: Non-existent File")
    print(f"{'=' * 60}")
    if not result:
        print("  ✅ PASS — Trả về empty string cho file không tồn tại")
    else:
        print("  ❌ FAIL — Không handle file không tồn tại")
    return not result


# ──────────────────────────────────────────────────────────────
# Main runner
# ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("🧪 PARSER FORMAT TEST SUITE")
    print("=" * 60)
    print("Testing parse_document() from src.utils.parser")
    print("Supported formats theo code: .pdf, .docx, .doc, .txt")
    print()

    results = {
        "PDF (.pdf)": test_pdf_parsing(),
        "DOCX (.docx)": test_docx_parsing(),
        "TXT UTF-8 (.txt)": test_txt_utf8_parsing(),
        "TXT Latin-1 (.txt)": test_txt_latin1_parsing(),
        "DOC cũ (.doc)": test_doc_old_format(),
        "Unsupported exts": test_unsupported_extension(),
    }
    test_empty_file()
    results["Non-existent file"] = test_nonexistent_file()

    print(f"\n{'=' * 60}")
    print("📊 TỔNG KẾT")
    print(f"{'=' * 60}")
    for name, passed in results.items():
        icon = "✅" if passed else "❌"
        print(f"  {icon} {name}")

    supported = sum(1 for v in results.values() if v)
    total = len(results)
    print(f"\n  Kết quả: {supported}/{total} tests passed")
