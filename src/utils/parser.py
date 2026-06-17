import os

import docx
import pdfplumber


def extract_text_from_pdf(file_path: str) -> str:
    """Trích xuất văn bản thô từ file PDF sử dụng pdfplumber."""
    text_content = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        return "\n\n".join(text_content)
    except Exception as e:
        print(f"Lỗi khi parse PDF {file_path}: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Trích xuất văn bản từ file DOCX sử dụng python-docx."""
    text_content = []
    try:
        doc = docx.Document(file_path)
        # 1. Trích xuất từ các đoạn văn (paragraphs)
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        # 2. Trích xuất từ các bảng biểu (tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))

        return "\n\n".join(text_content)
    except Exception as e:
        print(f"Lỗi khi parse DOCX {file_path}: {e}")
        return ""


def parse_document(file_path: str) -> str:
    """Hàm wrapper tự động nhận diện đuôi file và parse phù hợp."""
    if not os.path.exists(file_path):
        return ""

    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        try:
            with open(file_path, encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, encoding="latin-1") as f:
                return f.read()
    else:
        return ""


def safe_parse_bloom_level(bloom_val, default: int = 3) -> int:
    """Safely converts bloom_level string or other types to an integer between 1 and 6."""
    if bloom_val is None:
        return default
    if isinstance(bloom_val, int):
        if 1 <= bloom_val <= 6:
            return bloom_val
        return default

    # Try converting to string and extract first digit
    val_str = str(bloom_val).strip()
    import re

    match = re.search(r"\d+", val_str)
    if match:
        try:
            val_int = int(match.group(0))
            if 1 <= val_int <= 6:
                return val_int
        except ValueError:
            pass
    return default
